#####################################################################################
# The script is used to generate word document contains lookup table automatically. #
# The data is collected from .txt file in buflib folder								#
#####################################################################################

from docx import Document
from docx.shared import Inches,Pt

buflib = ['X1','X2','X3','X4','X5']

buffer_size = 	  {'X1':[2,1.37],'X2':[4,2.74],
				   'X3':[8,5.48],'X4':[16,10.96],
				   'X5':[32,21.92]}
load_cap = [15,20,25,30,35,40,45]
slew_in = [50,55,60,65,70,75,80]

document = Document()

document.add_heading('PTM45nm 0.55v下buffer查找表', 0)
style = document.styles['Normal']
font = style.font
font.name = 'SimHei'
font.size = Pt(8)

document.add_paragraph('时间单位：ps，电容单位：fF，功耗单位：uW')

document.add_paragraph('buffer_type  pmos(um) nmos(um)')
for key,value in buffer_size.items():
	document.add_paragraph('{}               {}    {}'.format(key,value[0],value[1]))

document.add_page_break()

for size in buflib:
	lut_fall = './buflib/{}/lut_fall.txt'.format(size)
	lut_rise = './buflib/{}/lut_rise.txt'.format(size)

	delay_table = [[] for i in range(7)]
	slew_table  = [[] for i in range(7)]
	power_table = [[] for i in range(7)]

	# update lut data
	with open(lut_rise) as f:
		f.readline()
		for i in range(7):
			for j in range(7):
				data = f.readline().split()
				delay_table[i].append((data[0],data[1]))
				slew_table[i].append((data[2],data[3]))
				power_table[i].append(data[4])


	document.add_heading(size, level=1)
	document.add_paragraph('输入信号上升沿', style='Intense Quote')
	document.add_paragraph('延迟及其波动表(均值/方差)', style='List Bullet')
	table = document.add_table(rows=1, cols=8)
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = "slew\\cap"
	for i,name in enumerate(load_cap):
		hdr_cells[i+1].text = str(name)
	for i,data in enumerate(delay_table):
		row_cells = table.add_row().cells
		row_cells[0].text = str(slew_in[i])
		for j, delay in enumerate(data):
			miu,sigma = delay
			row_cells[j+1].text = "{:.1f}/{:.2f}".format(float(miu)*1e12,float(sigma)*1e12)

	document.add_paragraph('slew及其波动表(均值/方差)', style='List Bullet')
	table = document.add_table(rows=1, cols=8)
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = "slew\\cap"
	for i,name in enumerate(load_cap):
		hdr_cells[i+1].text = str(name)
	for i,data in enumerate(slew_table):
		row_cells = table.add_row().cells
		row_cells[0].text = str(slew_in[i])
		for j, slew in enumerate(data):
			miu,sigma = slew
			row_cells[j+1].text = "{:.1f}/{:.2f}".format(float(miu)*1e12,float(sigma)*1e12)

	document.add_page_break()

	# update lut data
	delay_table = [[] for i in range(7)]
	slew_table  = [[] for i in range(7)]
	power_table = [[] for i in range(7)]
	with open(lut_fall) as f:
		f.readline()
		for i in range(7):
			for j in range(7):
				data = f.readline().split()
				delay_table[i].append((data[0],data[1]))
				slew_table[i].append((data[2],data[3]))
				power_table[i].append(data[4])

	document.add_heading(size, level=1)
	document.add_paragraph('输入信号下降沿', style='Intense Quote')
	document.add_paragraph('延迟及其波动表(均值/方差)', style='List Bullet')
	table = document.add_table(rows=1, cols=8)
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = "slew\\cap"
	for i,name in enumerate(load_cap):
		hdr_cells[i+1].text = str(name)
	for i,data in enumerate(delay_table):
		row_cells = table.add_row().cells
		row_cells[0].text = str(slew_in[i])
		for j, delay in enumerate(data):
			miu,sigma = delay
			row_cells[j+1].text = "{:.1f}/{:.2f}".format(float(miu)*1e12,float(sigma)*1e12)

	document.add_paragraph('slew及其波动表(均值/方差)', style='List Bullet')
	table = document.add_table(rows=1, cols=8)
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = "slew\\cap"
	for i,name in enumerate(load_cap):
		hdr_cells[i+1].text = str(name)
	for i,data in enumerate(slew_table):
		row_cells = table.add_row().cells
		row_cells[0].text = str(slew_in[i])
		for j, slew in enumerate(data):
			miu,sigma = slew
			row_cells[j+1].text = "{:.1f}/{:.2f}".format(float(miu)*1e12,float(sigma)*1e12)

	document.add_page_break()

document.save('buffer_lut.docx')

